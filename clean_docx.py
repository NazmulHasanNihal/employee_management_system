import os
from docx import Document

def clean_evm_terminology(file_path, output_path):
    doc = Document(file_path)
    
    replacements = {
        "EVM": "EMS",
        "evm": "ems",
        "Ethereum Virtual Machine": "Enterprise Management System",
        "Ethereum": "Enterprise",
        "Smart Contract": "Business Logic",
        "smart contract": "business logic",
        "blockchain": "ledger",
        "Blockchain": "Ledger",
        "Web3": "Web2.5",
        "web3": "web2.5"
    }

    count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
                    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for old_text, new_text in replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
                                
    doc.save(output_path)
    print(f"Successfully replaced {count} instances of EVM terminology in {file_path}, saved to {output_path}")

if __name__ == "__main__":
    docx_path = os.path.join("docs", "Project Report.docx")
    output_path = os.path.join("docs", "Project Report Cleaned.docx")
    clean_evm_terminology(docx_path, output_path)
