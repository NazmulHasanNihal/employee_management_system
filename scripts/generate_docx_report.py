import os
import sys

def create_report():
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Installing python-docx...")
        os.system(f"{sys.executable} -m pip install python-docx")
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    
    # Title
    title = doc.add_heading('OpsHub Enterprise EMS - Final Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Prepared by: OpsHub Core Engineering Team')
    doc.add_paragraph('Date: August 2026')
    doc.add_paragraph('CEO: Nazmul Hasan Nihal')
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The OpsHub Enterprise Employee Management System (EMS) project has been successfully completed, "
        "fully localized for the Bangladeshi market. The system now perfectly handles enterprise-grade "
        "HR operations, shift scheduling, automated payroll processing in Bangladeshi Taka (BDT), and "
        "asset management."
    )
    
    doc.add_heading('2. System Architecture & Tech Stack', level=1)
    doc.add_paragraph("Frontend: Next.js 15 (App Router), React, Tailwind CSS, Base-UI")
    doc.add_paragraph("Backend: Node.js / Next.js Server Actions, tRPC")
    doc.add_paragraph("Database: PostgreSQL via Prisma ORM")
    doc.add_paragraph("Authentication: Custom JWT + Supabase Auth integration")
    
    doc.add_heading('3. Localization & Data Realism', level=1)
    p = doc.add_paragraph()
    p.add_run("Currency Updates: ").bold = True
    p.add_run("All financial metrics, pricing tiers, ROI calculators, and asset book values have been strictly converted to Bangladeshi Taka (BDT), using accurate local enterprise metrics.")
    
    p = doc.add_paragraph()
    p.add_run("Branch Consolidation: ").bold = True
    p.add_run("The database architecture was streamlined to a single 'Dhaka HQ' branch to better reflect the current organizational structure.")
    
    p = doc.add_paragraph()
    p.add_run("Salary Normalization: ").bold = True
    p.add_run("All 100 seeded employees' base salaries have been adjusted to reflect realistic local market conditions (e.g., CEO at ৳5,00,000, Engineers at ৳30,000 - ৳1,00,000).")
    
    p = doc.add_paragraph()
    p.add_run("Hardware Asset Realism: ").bold = True
    p.add_run("Replaced synthetic hardware lists with real-world enterprise devices like HP EliteBook 840 G8, HP Elite x3 smartphones, and Dell Latitude series, accompanied by true-to-market purchase prices in BDT.")
    
    doc.add_heading('4. Security & Performance Fixes', level=1)
    doc.add_paragraph(
        "- Removed console.log statements and debug traces from critical production paths.\n"
        "- Hardened environment variable checks (e.g., DEMO_PASSWORD).\n"
        "- Resolved syntax layout errors (e.g., nested <T> tags in <option> elements in the Attendance and Payment hubs).\n"
        "- Implemented strict React-based client-side pagination to eliminate vertical scrolling overflow in data-heavy hubs."
    )
    
    doc.add_heading('5. Landing Page Optimization', level=1)
    doc.add_paragraph(
        "The SaaS landing page was completely rebuilt to showcase the true capabilities of OpsHub. "
        "Testimonials were updated to feature the CEO, Nazmul Hasan Nihal, alongside other localized "
        "enterprise use-cases. The Pricing and ROI calculator now effectively demonstrate value using BDT scaling."
    )
    
    # Save the file to artifacts as requested by user or save it in the root dir
    # I will save it in the project root directory where the python script expects it
    file_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'OpsHub_Enterprise_EMS_Project_Report.docx')
    doc.save(file_path)
    
    # Also save a copy to the artifacts directory so the AI can link to it
    artifact_path = r"C:\Users\Nazmul Hasan Nihal\.gemini\antigravity-ide\brain\10491cd9-79a5-4a47-88ee-ac5db4201942\OpsHub_Enterprise_EMS_Project_Report.docx"
    doc.save(artifact_path)
    
    print(f"Report successfully generated at: {file_path}")

if __name__ == "__main__":
    create_report()
